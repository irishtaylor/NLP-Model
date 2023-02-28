import string 
import language_tool_python  # library for checking grammar - pip install language-tool-python
from textblob import TextBlob # python library to check spellings and grammar of a text - pip install textblob
import contractions # python library to fix contractions words from a paragraph text  - pip install contractions
from nltk.corpus import stopwords
import nltk 
import csv 
import spacy #NLP library - pip install spacy
from docx.shared import Inches 
from spacy.language import Language
from spacy_langdetect import LanguageDetector # spacy library - function to detect language of a text
import pandas
from pandas import *
import docx # library to auto-genearte word doc
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from docx2pdf import convert

name_test = "Aakriti"  # need to make candidate name dynamic

#val = input("Enter your input paragraph: ")

def get_lang_detector(nlp, name):  # function to initiate and return LanguageDetector() fun for NLLP 
    return LanguageDetector()

def percentageOfBad(x):   # function to caalculate percentage of baad spellings comapred to the good
    return (x[0] / (x[1] + x[0])) * 100

def compare(text1, text2):  
    l1 = text1.split()  
    l2 = text2.split()
    good = 0
    bad = 0
    for i in range(0, len(l1)):
        if l1[i] != l2[i]:
            bad += 1
        else:
            good += 1
    return (good, bad)

def evaluate (words_to_evaluate):  # main function doing all evaluation of checking spellings and grammaar of the candidate essay text
    words_to_evaluate=input("Enter the essay:")   # candidate essay input text

    #1st Grammar and spelling checking layer
    # using the langaue-tool-python to check grammar errors in a text  
    my_tool = language_tool_python.LanguageTool('en-US', config={'maxSpellingSuggestions': 1}) # initiatiing the library language-tool-python 
    # will need JAVA 8 to be installed to work with this language-tool-python library to checking grammar errors in the text

      # to count words in string
    res = sum([i.strip(string.punctuation).isalpha() for i in words_to_evaluate.split()])   

    no_of_words_in_essay = str(res)

    # printing result
    print ("The number of words are : " +  str(res))

    from spellchecker import SpellChecker 

    spell = SpellChecker()  # initiating SpellChecker library

    #wordlist=strp.split()
    
    import re
    strp=re.findall("[a-zA-Z""]+",words_to_evaluate) # text cleaning -important step
    updated_docx=(" ".join(strp))
    contractions_corrected_text = contractions.fix(words_to_evaluate, slang=True)  # fixing contractions in the text
    print(contractions_corrected_text)
    # importing module
    misspelled = list(spell.unknown(strp)) # detecting misspelled words in the text
    no_of_spellings_count = len(misspelled)

    # Uncomment these 2 lines to download the language specific libraries related with Spacy libraary at the first time to perform further NLP prcessing of the essay 
    # and then comment back to continue
    
    # PLEASE BE PATIENT AT THIS TEP to download spacy cli
    # this step - of initial download of spacy cli download and load will take around 5-10 mins , screenshot of successfull installation given along with code
    
    # after first time code run - spacy-cli will be downloaded
    # uncomment the below lines after that
    
    #spacy.cli.download("en_core_web_sm")
    #spacy.cli.download("en")
    text = words_to_evaluate

    nlp = spacy.load("en_core_web_sm") # loading spacy library for any NLP processing - important step
    Language.factory("language_detector", func=get_lang_detector)
    nlp.add_pipe('language_detector', last=True) # initiating NLP - spacy pipeline - for detecting laanguage of the text
    # this functionality can be elaborated to work on
    doc = nlp(text)

    #print(doc._.language)

    nytimes= nlp(words_to_evaluate)
    entities=[(i, i.label_, i.label) for i in nytimes.ents] # detectoing Special, proper nouns, emtities in the text

    entities_list = []

    #print(entities)   
    for ent in entities:   
        entities_list.append(ent[0])   # appending all the entities text in a sperate list - to later detect them from misspelled errors list to avoid them to be flagged as errors
        print(ent[0])
    print("this is entities list", entities_list)

    # reading CSV file
    data = read_csv("words_new.csv")   # reading the Oxford/English dictionary words dataset (of 6,000,00 or more enteries) to detect them from misspelled errors list to avoid them to be flagged as errors
    # converting column data to list

    fc = data['words'].tolist()   # converting the English dictionary words dataset to a list
    test_list = misspelled
    
    # initializing remove list
    remove_list = fc

    # printing original list
    #print ("The original list is : " + str(test_list))

    # printing remove list
    #print ("The removal dataset list is : " + str(remove_list))

    if not set(test_list).isdisjoint(set(remove_list)):  # checking for duplicate errors /words in the misspelled errors list and removing them - for unique errors
        print("Duplicates found.")
        # using list comprehension to perform task
        res = [i for i in test_list if i not in remove_list]

        # printing result
        print ("The list after performing remove operation is : " + str(res))
    else:

        print("No duplicates found.")
    
    

    if not set(test_list).isdisjoint(set(entities_list)): # checking for entities words in the misspelled errors list and removing them - to avoid them to be flagged as erors in final list
        print("Proper nouns Duplicates found.")
        # using list comprehension to perform task
        res = [i for i in test_list if i not in entities_list]

        # printing result
        print ("The list after performing remove operation is : " + str(res))
    else:
        print("No duplicates found.")

    print("updated grammar errors list", test_list) # final updated grammar  erorrs list

    # 2nd Grammar and spelling checking layer - using Textblob library 
    spellings_checking = TextBlob(words_to_evaluate)

    correct_spellings = spellings_checking.correct() # genearting corrected text 
    
    #print("test", spellings_checking)
    mistakesCompCorrected = compare(words_to_evaluate, correct_spellings)  # comparing corercted text with the original candidate text

    print("MISSPELLING",len(res),res)
    print("Spelling Mistakes compared to corrected ", mistakesCompCorrected, "\n")

    print("Percentage of fixed spelling mistakes: ", percentageOfBad(mistakesCompCorrected), "%", "\n")

    
    #tet = [spell.correction(word) for word in words_to_evaluate]

    # count Word frequencies

    betty = TextBlob(words_to_evaluate)
    word_frequencies_for_word_event = betty.word_counts['event']
    #print("Checking the word frequency count for word - Event ", word_frequencies)

    word_frequencies_life = betty.word_counts['life']
    #print("Checking the word frequency count for word - my Life ", word_frequencies)

    
    # = Word(words_to_evaluate)
    #print("teststs", w.spellcheck())
    # getting the matches  

    # block for grammar checkking - language-tool-python

    #initiating some variables , for mistakes , corrections errors fpr grammar checking
    my_matches = my_tool.check(words_to_evaluate)  
    
    # defining some variables  
    myMistakes = []  
    myCorrections = []  
    startPositions = []  
    endPositions = []  
    
    # using the for-loop  
    for rules in my_matches:   # defining rules for grammar check 
        
        if len(rules.replacements) > 0:  
            startPositions.append(rules.offset)  
            endPositions.append(rules.errorLength + rules.offset)  
            myMistakes.append(words_to_evaluate[rules.offset : rules.errorLength + rules.offset])  
            myCorrections.append(rules.replacements[0])  
            
    # creating new object  
    my_NewText = list(words_to_evaluate)   
    
    # rewriting the correct passage  
    for n in range(len(startPositions)):  
        for i in range(len(words_to_evaluate)):  
            my_NewText[startPositions[n]] = myCorrections[n]  
            if (i > startPositions[n] and i < endPositions[n]):  
                my_NewText[i] = ""  
    
    my_NewText = "".join(my_NewText) 

    
    print('mistakes of the essay are', len(myMistakes), myMistakes)  
    
    #corrections_set = set(myCorrections)
    #print(list(corrections_set))
    #print("corrections are:", myCorrections)

    # correction  
    correct_text = my_tool.correct(my_NewText)  

    #print(correct_text)
    ##underlined_text = "\x1B[4m" + correct_text + "\x1B[0m"
    
    for x in res:
        myMistakes.append(x)

    print("comprised list ", len(myMistakes), myMistakes)
    

    # final unique merged and filtred list of grammar and spelling errrors taken for both the layers
    mistakes_set = set(myMistakes) # filtering duplicate errors out of initial grammar errors list
    final_list = list(mistakes_set)
    print("this is unique list", len(final_list), final_list)
    
    
    doc = nlp(words_to_evaluate)
    #print("One set of Keywords extracted using spacy library are " , doc.ents)
    
    #displacy.render(nytimes, style = "ent")

    #code for auto-generating word document

    # Create an instance of a word document
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run('\t\t\t\t\t')
    r.add_picture('MicrosoftTeams-image.png')
    
    #doc.add_picture('og_healthcare.full.colour.ahi.jpg',width=Inches(1.0), height=Inches(.2))
    #doc.add_run('\t\t\t\t\t')
    # Add a Title to the document 
    #doc.add_heading('\t\t\t\t Essay Topic', 0)

    # Adding paragraph with Increased font size

    # Add black title
    styles = doc.styles
    styles['Heading 3'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.name = "Times New Roman"

    #styles['Heading3'].font.name = 'ms sans serif'
    #font.size = docx.shared.Pt(16)

    #ont.size = docx.shared.Pt(12)
    #doc.add_heading('Title', level=2)

    from docx.enum.style import WD_STYLE_TYPE

    styles = doc.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 3']
    font = new_heading_style.font
    font.name = 'ms sans serif'
    font.size = Pt(14)
    doc.add_paragraph('Recruit Id : 003', style='New Heading')
    doc.add_paragraph('Candidate Name', style='New Heading')
    doc.add_paragraph('Essay Topic', style='New Heading')
    doc.add_heading("", 3)

    
    #header = doc.add_heading('Recruit Id : 001', 3)
    #header.style.font.name = 'Arial'
    #eader.style.font.size = Pt(14)
    
    #doc.add_heading('Candidate Name : ' +  str(name_test), 3)
    #doc.add_heading('Candidate Name : ' , 3)
    #doc.add_heading("Essay Topic : ", 2)
    #doc.add_heading("", 3)
    r.add_break()
    r.font.color.rgb = RGBColor(0,0,0)
    #r.add_break()
    #doc.add_heading("Essay written by the candidate(which has errors)", 3)

    '''Apply style'''
    style = doc.styles['Normal']
    font = style.font
    font.name = 'ms sans serif'
    font.size = docx.shared.Pt(12)
    para = doc.add_paragraph().add_run(words_to_evaluate)
    para.alignment = 3
    #doc.add_run('Green ')

    #wp = p.add_run('I want this sentence colored red with fontsize=22')
    #wp.font.size = Pt(22)
    para.font.color.rgb = RGBColor(0,0,0)
    
    # Table data in a form of list
    data = (
        ('Word Count : ', no_of_words_in_essay ),
        ('Grammar and spelling errors count : ', len(final_list)),
        ('Grammar and spelling errors list : ', final_list),
        ('Frequency for keyword : <Event>' , word_frequencies_for_word_event),
        ('Frequency for keyword : <Life>', word_frequencies_life),
        ('Special words or entities : ', entities)
    )
    
    # Creating a table object
    table = doc.add_table(rows=1, cols=2)
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    run = row[0].paragraphs[0].add_run('Labels')
    run.bold = True
    run = row[1].paragraphs[0].add_run('Data')
    run.bold = True
    
    # Adding data from the list to the table
    for id, name in data:
    
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = str(name)

    # Increasing size of the font
    para.font.size = Pt(12)

    # Choosing the top most section of the page
    section = doc.sections[0]
    
    # Calling the footer
    footer = section.footer
    
    # Calling the paragraph already present in
    # the footer section
    footer_para = footer.paragraphs[0]
    #footer_para.add_run('\t\t\t\t')

    #footer 
    # Adding the left zoned footer
    footer_para.text = " ©  All Rights Reserved | OG Healthcare | Essay Evaluation Report"
    

    #doc.add_paragraph(' '.join(list(myset)))
    #doc.add_paragraph(*myset, end=",")

    # Now save the document to a location 
    doc.save('auto-generated-docs/Id_003_name_date.docx')


    # as the python file
    convert("auto-generated-docs/Id_003_name_date.docx")
    
    # Converting docx specifying both the input
    # and output paths
    convert("auto-generated-docs/Id_003_name_date.docx", "auto-generated-docs/Id_003_name_date.pdf")

words_to_evaluate=1   #any default value
evaluate(words_to_evaluate)